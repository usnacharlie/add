from flask import Blueprint, render_template, request, redirect, url_for, flash, session, jsonify, send_file
from werkzeug.utils import secure_filename
import pandas as pd
import docx
import os
import re
from datetime import datetime
import requests
import io
import csv

import_bp = Blueprint('import_data', __name__, url_prefix='/import')

# Backend API URL
API_URL = "http://localhost:57021/api/v1"

# Allowed file extensions
ALLOWED_EXTENSIONS = {'csv', 'xlsx', 'xls', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def admin_required(f):
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_token' not in session or session.get('user_role') != 'admin':
            flash('Admin privileges required.', 'danger')
            return redirect(url_for('auth.login'))
        return f(*args, **kwargs)
    return decorated_function

def validate_nrc(nrc):
    """Validate NRC format: 123456/78/9"""
    pattern = r'^\d{6}/\d{2}/\d{1}$'
    return bool(re.match(pattern, nrc))

def validate_phone(phone):
    """Validate Zambian phone number"""
    # Remove any spaces or special characters
    phone = re.sub(r'[^\d]', '', phone)
    # Check if it's a valid Zambian number
    if len(phone) == 10 and phone.startswith('0'):
        return phone
    elif len(phone) == 12 and phone.startswith('260'):
        return '0' + phone[3:]
    return None

def parse_gender(gender_str):
    """Parse gender from various formats"""
    if not gender_str:
        return None
    gender_str = str(gender_str).strip().upper()
    if gender_str in ['M', 'MALE']:
        return 'Male'
    elif gender_str in ['F', 'FEMALE']:
        return 'Female'
    return 'Other'

@import_bp.route('/')
@admin_required
def index():
    """Main import page"""
    import_history = []

    # Get import history from backend
    try:
        headers = {'Authorization': f'Bearer {session.get("user_token")}'}
        response = requests.get(f"{API_URL}/imports/history", headers=headers)
        if response.status_code == 200:
            import_history = response.json()
    except:
        pass

    return render_template('import/index.html', history=import_history)

@import_bp.route('/upload', methods=['GET', 'POST'])
@admin_required
def upload():
    """Upload and process bulk member data"""
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file selected', 'danger')
            return redirect(url_for('import_data.upload'))

        file = request.files['file']
        if file.filename == '':
            flash('No file selected', 'danger')
            return redirect(url_for('import_data.upload'))

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_ext = filename.rsplit('.', 1)[1].lower()

            # Save file temporarily
            temp_path = os.path.join('/tmp', filename)
            file.save(temp_path)

            try:
                # Process based on file type
                if file_ext in ['csv']:
                    data = process_csv(temp_path)
                elif file_ext in ['xlsx', 'xls']:
                    data = process_excel(temp_path)
                elif file_ext == 'docx':
                    data = process_docx(temp_path)
                else:
                    flash('Unsupported file type', 'danger')
                    return redirect(url_for('import_data.upload'))

                # Store in session for preview
                session['import_data'] = data
                session['import_filename'] = filename

                # Clean up temp file
                os.remove(temp_path)

                return redirect(url_for('import_data.preview'))

            except Exception as e:
                flash(f'Error processing file: {str(e)}', 'danger')
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                return redirect(url_for('import_data.upload'))
        else:
            flash('Invalid file type. Allowed types: CSV, Excel, DOCX', 'danger')
            return redirect(url_for('import_data.upload'))

    return render_template('import/upload.html')

def process_csv(filepath):
    """Process CSV file and extract member data"""
    members = []

    with open(filepath, 'r', encoding='utf-8-sig') as file:
        # Try to detect delimiter
        sample = file.read(1024)
        file.seek(0)
        sniffer = csv.Sniffer()
        delimiter = sniffer.sniff(sample).delimiter

        reader = csv.DictReader(file, delimiter=delimiter)

        for row in reader:
            member = extract_member_from_row(row)
            if member:
                members.append(member)

    return members

def process_excel(filepath):
    """Process Excel file and extract member data"""
    members = []

    # Read Excel file
    df = pd.read_excel(filepath)

    # Convert to dictionary records
    for _, row in df.iterrows():
        member = extract_member_from_row(row.to_dict())
        if member:
            members.append(member)

    return members

def process_docx(filepath):
    """Process DOCX file and extract member data from tables"""
    members = []

    doc = docx.Document(filepath)

    # Extract location info from header
    location_info = {'province': '', 'district': '', 'ward': ''}
    for paragraph in doc.paragraphs:
        text = paragraph.text.upper()
        if 'PROVINCE:' in text:
            # Try to extract province, district, ward
            parts = text.split()
            for i, part in enumerate(parts):
                if 'PROVINCE:' in part and i+1 < len(parts):
                    location_info['province'] = parts[i+1].strip('.')
                elif 'DISTRICT:' in part and i+1 < len(parts):
                    location_info['district'] = parts[i+1].strip('.')
                elif 'WARD:' in part and i+1 < len(parts):
                    location_info['ward'] = parts[i+1].strip('.')

    # Process tables
    for table in doc.tables:
        # Skip header row
        for i, row in enumerate(table.rows):
            if i == 0:
                continue  # Skip header

            cells = [cell.text.strip() for cell in row.cells]

            # Check if row has data (not just row number)
            if len(cells) >= 6 and cells[1]:  # Name field is not empty
                member = {
                    'name': cells[1],
                    'gender': parse_gender(cells[2]),
                    'age': cells[3],
                    'nrc_number': cells[4],
                    'voter_id': cells[5],
                    'phone': validate_phone(cells[6]) if len(cells) > 6 else '',
                    'province': location_info['province'],
                    'district': location_info['district'],
                    'ward': location_info['ward']
                }

                # Split name into first and last
                name_parts = member['name'].split()
                if len(name_parts) >= 2:
                    member['first_name'] = name_parts[0]
                    member['last_name'] = ' '.join(name_parts[1:])
                else:
                    member['first_name'] = member['name']
                    member['last_name'] = ''

                # Validate and clean data
                if member['nrc_number'] and validate_nrc(member['nrc_number']):
                    members.append(member)

    return members

def extract_member_from_row(row):
    """Extract member data from a row (dict)"""
    member = {}

    # Map common column names
    name_fields = ['name', 'Name', 'NAME', 'full_name', 'Full Name']
    fname_fields = ['first_name', 'First Name', 'FIRST_NAME', 'fname']
    lname_fields = ['last_name', 'Last Name', 'LAST_NAME', 'lname', 'surname']
    nrc_fields = ['nrc', 'NRC', 'nrc_number', 'NRC Number', 'national_id']
    phone_fields = ['phone', 'Phone', 'PHONE', 'contact', 'Contact', 'mobile']
    gender_fields = ['gender', 'Gender', 'GENDER', 'sex', 'Sex']
    age_fields = ['age', 'Age', 'AGE']
    voter_fields = ['voter_id', 'Voter ID', "Voter's ID", 'voter_registration']

    # Extract name
    for field in name_fields:
        if field in row and row[field]:
            name_parts = str(row[field]).split()
            if len(name_parts) >= 2:
                member['first_name'] = name_parts[0]
                member['last_name'] = ' '.join(name_parts[1:])
            break

    # Or extract first and last name separately
    for field in fname_fields:
        if field in row and row[field]:
            member['first_name'] = str(row[field]).strip()
            break

    for field in lname_fields:
        if field in row and row[field]:
            member['last_name'] = str(row[field]).strip()
            break

    # Extract NRC
    for field in nrc_fields:
        if field in row and row[field]:
            nrc = str(row[field]).strip()
            if validate_nrc(nrc):
                member['nrc_number'] = nrc
            break

    # Extract phone
    for field in phone_fields:
        if field in row and row[field]:
            phone = validate_phone(str(row[field]))
            if phone:
                member['phone'] = phone
            break

    # Extract gender
    for field in gender_fields:
        if field in row and row[field]:
            member['gender'] = parse_gender(row[field])
            break

    # Extract age
    for field in age_fields:
        if field in row and row[field]:
            try:
                age = int(float(str(row[field])))
                if 18 <= age <= 120:
                    member['age'] = age
            except:
                pass
            break

    # Extract voter ID
    for field in voter_fields:
        if field in row and row[field]:
            member['voter_registration_number'] = str(row[field]).strip()
            break

    # Only return if we have minimum required fields
    if member.get('first_name') and member.get('nrc_number'):
        return member

    return None

@import_bp.route('/preview')
@admin_required
def preview():
    """Preview data before importing"""
    data = session.get('import_data', [])
    filename = session.get('import_filename', 'Unknown')

    if not data:
        flash('No data to preview', 'warning')
        return redirect(url_for('import_data.upload'))

    # Validate each record
    for member in data:
        member['valid'] = bool(
            member.get('first_name') and
            member.get('nrc_number') and
            validate_nrc(member.get('nrc_number', ''))
        )
        member['errors'] = []

        if not member.get('first_name'):
            member['errors'].append('Missing first name')
        if not member.get('last_name'):
            member['errors'].append('Missing last name')
        if not validate_nrc(member.get('nrc_number', '')):
            member['errors'].append('Invalid NRC format')
        if member.get('phone') and not validate_phone(member.get('phone')):
            member['errors'].append('Invalid phone number')

    valid_count = sum(1 for m in data if m['valid'])
    invalid_count = len(data) - valid_count

    return render_template('import/preview.html',
                         data=data,
                         filename=filename,
                         total=len(data),
                         valid=valid_count,
                         invalid=invalid_count)

@import_bp.route('/process', methods=['POST'])
@admin_required
def process():
    """Process the import"""
    data = session.get('import_data', [])

    if not data:
        flash('No data to import', 'warning')
        return redirect(url_for('import_data.upload'))

    # Filter valid records only
    valid_members = [m for m in data if m.get('valid', False)]

    if not valid_members:
        flash('No valid records to import', 'danger')
        return redirect(url_for('import_data.preview'))

    # Prepare for import
    import_batch = {
        'filename': session.get('import_filename', 'Unknown'),
        'total_records': len(data),
        'valid_records': len(valid_members),
        'imported_by': session.get('user_id'),
        'members': []
    }

    success_count = 0
    error_count = 0
    errors = []

    # Process each member
    for member in valid_members:
        # Prepare member data for API
        member_data = {
            'first_name': member.get('first_name'),
            'last_name': member.get('last_name', ''),
            'nrc_number': member.get('nrc_number'),
            'phone': member.get('phone', ''),
            'gender': member.get('gender', 'Other'),
            'age': member.get('age'),
            'voter_registration_number': member.get('voter_registration_number', ''),
            'password': member.get('nrc_number', '').replace('/', ''),  # Default password is NRC without slashes
            'registration_channel': 'bulk_import'
        }

        # Add location if available
        if member.get('province'):
            member_data['province'] = member['province']
        if member.get('district'):
            member_data['district'] = member['district']
        if member.get('ward'):
            member_data['ward'] = member['ward']

        # Send to API
        try:
            headers = {'Authorization': f'Bearer {session.get("user_token")}'}
            response = requests.post(f"{API_URL}/members/register",
                                    json=member_data,
                                    headers=headers)

            if response.status_code in [200, 201]:
                success_count += 1
                result = response.json()
                import_batch['members'].append({
                    'name': f"{member['first_name']} {member.get('last_name', '')}",
                    'membership_number': result.get('membership_number'),
                    'status': 'success'
                })
            else:
                error_count += 1
                error_msg = response.json().get('detail', 'Import failed')
                errors.append(f"{member['first_name']} {member.get('last_name', '')}: {error_msg}")

        except Exception as e:
            error_count += 1
            errors.append(f"{member['first_name']} {member.get('last_name', '')}: {str(e)}")

    # Clear session data
    session.pop('import_data', None)
    session.pop('import_filename', None)

    # Show results
    flash(f'Import completed: {success_count} successful, {error_count} failed',
          'success' if error_count == 0 else 'warning')

    if errors:
        for error in errors[:10]:  # Show first 10 errors
            flash(error, 'danger')

    return redirect(url_for('import_data.results',
                          success=success_count,
                          failed=error_count))

@import_bp.route('/results')
@admin_required
def results():
    """Show import results"""
    success = request.args.get('success', 0, type=int)
    failed = request.args.get('failed', 0, type=int)

    return render_template('import/results.html',
                         success=success,
                         failed=failed,
                         total=success+failed)

@import_bp.route('/template/<format>')
@admin_required
def download_template(format):
    """Download import template"""
    if format == 'csv':
        # Create CSV template
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(['First Name', 'Last Name', 'NRC Number', 'Phone', 'Gender', 'Age',
                        'Voter ID', 'Province', 'District', 'Ward', 'Email', 'Occupation'])
        writer.writerow(['John', 'Doe', '123456/78/9', '0977123456', 'M', '25',
                        '12345678', 'Lusaka', 'Lusaka', 'Kabwata', 'john@example.com', 'Teacher'])

        output.seek(0)
        return send_file(
            io.BytesIO(output.getvalue().encode()),
            mimetype='text/csv',
            as_attachment=True,
            download_name='member_import_template.csv'
        )

    elif format == 'excel':
        # Create Excel template
        df = pd.DataFrame({
            'First Name': ['John', ''],
            'Last Name': ['Doe', ''],
            'NRC Number': ['123456/78/9', ''],
            'Phone': ['0977123456', ''],
            'Gender': ['M', ''],
            'Age': [25, ''],
            'Voter ID': ['12345678', ''],
            'Province': ['Lusaka', ''],
            'District': ['Lusaka', ''],
            'Ward': ['Kabwata', ''],
            'Email': ['john@example.com', ''],
            'Occupation': ['Teacher', '']
        })

        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Members', index=False)

        output.seek(0)
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='member_import_template.xlsx'
        )

    else:
        flash('Invalid template format', 'danger')
        return redirect(url_for('import_data.index'))

@import_bp.route('/validate', methods=['POST'])
@admin_required
def validate_data():
    """AJAX endpoint to validate import data"""
    data = request.json.get('data', [])

    results = []
    for row in data:
        errors = []

        # Validate required fields
        if not row.get('first_name'):
            errors.append('First name required')
        if not row.get('nrc_number'):
            errors.append('NRC number required')
        elif not validate_nrc(row['nrc_number']):
            errors.append('Invalid NRC format')
        if row.get('phone') and not validate_phone(row['phone']):
            errors.append('Invalid phone number')

        results.append({
            'valid': len(errors) == 0,
            'errors': errors
        })

    return jsonify(results)